# -*- coding: utf-8 -*-
"""
audit_corriges.py — mesure la couverture des corrigés de TOUS les manuels VÉRITAS
et écrit AUDIT_CORRIGES.md (point sur l'existant + trous à combler).

Rien n'est déduit « de mémoire » : chaque chiffre est recompté depuis les fichiers
sources (docx des Guides pour le 1er cycle, sources balisées des packs pour le 2nd
cycle, cahiers EST, manuels d'œuvre).

Usage : python tools/audit_corriges.py
"""
import os, re, sys, io, glob, datetime
from collections import Counter, defaultdict

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
from docx import Document

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DESK = r"C:\Users\Mythe Errant\Desktop\Manuels"
FIN1 = os.path.join(DESK, "FINAUX_1er_cycle")
FIN2 = os.path.join(DESK, "FINAUX_2nd_cycle")
PACKS = {"2nde": "Manuel_2nde_A", "1ere": "Manuel_1ere_A", "tle": "Manuel_Tle_A"}
EST = os.path.join(ROOT, "Manuel_EST")
TAG = re.compile(r"^#([A-Z_0-9]+)::\s?(.*)$")

def nature(d):
    d = d or ""
    if "Sujet d'entraînement" in d or d.startswith("📘"):
        return "sujet d'entraînement (vierge par conception)"
    if "Banque de sujets" in d or d.startswith("📗"):
        return "banque de sujets"
    if d.startswith("⚖️") or "Épreuve" in d:
        return "épreuve / évaluation"
    if d.startswith("🩺"):
        return "évaluation diagnostique"
    if d.startswith("📖"):
        return "étude de l'œuvre intégrale"
    if d.startswith("📒"):
        return "groupement de textes"
    if d.startswith("🪶"):
        return "méthodologie / intégration"
    return "leçon"

# ── 1er cycle ───────────────────────────────────────────────────────────────
def audit_1er(niv):
    d = Document(os.path.join(FIN1, "%s_Guide_pedagogique.docx" % niv))
    seq = sem = disc = None
    exos, cur = [], None
    nseq = nsem = ndisc = 0
    for p in d.paragraphs:
        st = p.style.name if p.style is not None else ""
        t = p.text.strip()
        if not t:
            continue
        if st == "T_Seq":
            seq = t; nseq += 1
        elif st == "T_Sem":
            sem = t; nsem += 1
        elif st == "T_Disc":
            disc = t; ndisc += 1
        elif st == "Exercice":
            cur = dict(seq=seq, sem=sem, disc=disc, txt=t, corr=False, conseil=False)
            exos.append(cur)
        elif st == "Corrige" and cur is not None:
            cur["corr"] = True
        elif st == "Conseil" and cur is not None:
            cur["conseil"] = True
    sans = [e for e in exos if not e["corr"]]
    return dict(seq=nseq, sem=nsem, lecons=ndisc, exos=len(exos),
                corriges=len(exos) - len(sans), sans=sans)

# ── 2nd cycle ───────────────────────────────────────────────────────────────
def audit_2nd(niv):
    cdir = os.path.join(DESK, PACKS[niv], "pack", "content")
    blocks, cur, item = [], None, None
    seq = None
    for f in sorted(glob.glob(os.path.join(cdir, "*.md"))):
        for line in open(f, encoding="utf-8"):
            m = TAG.match(line.rstrip("\n"))
            if not m:
                continue
            tag, val = m.group(1), m.group(2).strip()
            if tag == "SEQ":
                seq = val.split("::")[0].strip()
            elif tag == "DISC":
                cur = dict(disc=val, seq=seq, file=os.path.basename(f), n=0, ok=0)
                blocks.append(cur); item = None
            elif tag in ("Q", "QW", "QN", "EXO"):
                if cur is None:
                    continue
                cur["n"] += 1; item = False
            elif tag in ("SOL", "SOLM"):
                if cur is None or item is None:
                    continue
                if item is False:
                    cur["ok"] += 1; item = True
    return blocks

def main():
    now = datetime.date.today().strftime("%d/%m/%Y")
    L = []
    W = L.append
    W("# Audit des corrigés des manuels VÉRITAS — %s" % now)
    W("")
    W("> Généré par `tools/audit_corriges.py`. Tous les chiffres sont recomptés depuis les fichiers")
    W("> sources (Guides `.docx` du 1er cycle, sources balisées des packs du 2nd cycle, cahiers EST).")
    W("")

    # 1er cycle
    W("## 1. Premier cycle (6ᵉ → 3ᵉ) — corrigés dans le **Guide pédagogique**")
    W("")
    W("| Niveau | Séquences | Semaines | Leçons | Exercices | Corrigés | Couverture | Sans corrigé |")
    W("|---|---:|---:|---:|---:|---:|---:|---|")
    manques1 = []
    for niv in ("6e", "5e", "4e", "3e"):
        r = audit_1er(niv)
        pct = 100.0 * r["corriges"] / max(1, r["exos"])
        detail = "—" if not r["sans"] else "%d (production orale/écrite : note enseignant)" % len(r["sans"])
        W("| %s | %d | %d | %d | %d | %d | **%.1f %%** | %s |"
          % (niv, r["seq"], r["sem"], r["lecons"], r["exos"], r["corriges"], pct, detail))
        for e in r["sans"]:
            manques1.append((niv, e))
    W("")
    sans_rien = [(n, e) for n, e in manques1 if not e["conseil"]]
    W("Exercices sans corrigé **ni** consigne d'évaluation pour l'enseignant : **%d**." % len(sans_rien))
    for n, e in sans_rien[:20]:
        W("- %s · %s · %s" % (n, (e["disc"] or "?")[:60], e["txt"][:90]))
    W("")

    # 2nd cycle
    W("## 2. Second cycle séries A (2ⁿᵈᵉ → Tˡᵉ) — corrigés dans le **Guide pédagogique** (balises `#SOL`)")
    W("")
    W("| Niveau | Blocs | Questions + exercices | Corrigés | Couverture |")
    W("|---|---:|---:|---:|---:|")
    trous = {}
    for niv in ("2nde", "1ere", "tle"):
        b = audit_2nd(niv)
        n = sum(x["n"] for x in b); ok = sum(x["ok"] for x in b)
        W("| %s | %d | %d | %d | **%.1f %%** |" % (niv, len(b), n, ok, 100.0 * ok / max(1, n)))
        trous[niv] = b
    W("")
    W("### Répartition des manques par nature de bloc")
    W("")
    W("| Niveau | Nature | Blocs | Questions | Corrigées | Manquantes |")
    W("|---|---|---:|---:|---:|---:|")
    for niv, b in trous.items():
        agg = defaultdict(lambda: [0, 0, 0])
        for x in b:
            a = agg[nature(x["disc"])]
            a[0] += 1; a[1] += x["n"]; a[2] += x["ok"]
        for k, (nb, q, ok) in sorted(agg.items(), key=lambda i: -(i[1][1] - i[1][2])):
            if q - ok == 0:
                continue
            W("| %s | %s | %d | %d | %d | **%d** |" % (niv, k, nb, q, ok, q - ok))
    W("")
    W("### Blocs à combler en priorité (≥ 3 questions, aucun corrigé)")
    W("")
    W("| Niveau | Fichier source | Questions | Bloc |")
    W("|---|---|---:|---|")
    for niv, b in trous.items():
        for x in b:
            if x["n"] >= 3 and x["ok"] == 0 and "entraînement" not in nature(x["disc"]):
                W("| %s | `%s` | %d | %s |" % (niv, x["file"], x["n"], x["disc"][:80]))
    W("")

    # EST
    W("## 3. Cahiers EST (séries scientifiques & techniques)")
    W("")
    W("| Cahier | Questions + exercices | Réponses dans le cahier (`#REP`) | Corrigés publiés en ligne |")
    W("|---|---:|---:|---|")
    for niv, lab in (("2nde_ST", "2ⁿᵈᵉ S&T"), ("1ere_ST", "1ʳᵉ S&T"), ("Tle_ST", "Tˡᵉ S&T")):
        c = Counter()
        for f in glob.glob(os.path.join(EST, niv, "content", "*.md")):
            for line in open(f, encoding="utf-8"):
                m = TAG.match(line.rstrip("\n"))
                if m:
                    c[m.group(1)] += 1
        page = "corriges/est-%s.html" % niv.split("_")[0].lower().replace("tle", "tle")
        ok = "oui — `%s`" % page if os.path.exists(os.path.join(ROOT, page)) else "**non**"
        W("| %s | %d | %d | %s |" % (lab, c["Q"] + c["EXO"] + c["QW"] + c["QN"], c["REP"], ok))
    W("")
    W("Les cahiers EST sont **conçus sans corrigés imprimés** (astuce : « Les corrigés existent, mais ils")
    W("ne profitent qu'à celui qui a d'abord cherché »). Les corrigés en ligne couvrent les **évaluations**")
    W("et les **tâches d'intégration**, pas les exercices de leçon.")
    W("")

    # Manuels d'œuvre
    W("## 4. Manuels d'œuvre (études d'œuvres intégrales)")
    W("")
    W("| Manuel | Sections | Sujets entièrement rédigés |")
    W("|---|---|---|")
    for d, lab in (("Manuel_Stances", "Stances et Poèmes (Sully Prudhomme)"),
                   ("Manuel_Ngum", "Ngum a Jemea (Mbanga Eyombwan)"),
                   ("Manuel_VieuxNegre", "Le Vieux Nègre et la Médaille (Oyono)"),
                   ("Manuel_Balafon", "Balafon (Engelbert Mveng)")):
        files = sorted(os.path.basename(x) for x in glob.glob(os.path.join(ROOT, d, "content", "*.js")))
        mod = [f for f in files if "modele" in f]
        W("| %s | %d | %s |" % (lab, len(files), "oui — `%s`" % mod[0] if mod else "**aucune section de corrigés**"))
    W("")

    # Renvois imprimés
    W("## 5. Renvois imprimés (« astuces ») vers le web")
    W("")
    W("| Manuel | Renvoi préface (URL `manuels.html`) | Renvois fin de séquence |")
    W("|---|---|---:|")
    for niv in ("6e", "5e", "4e", "3e"):
        d = Document(os.path.join(FIN1, "%s_Bord_cahier_de_francais.docx" % niv))
        txt = [p.text for p in d.paragraphs]
        pref = sum(1 for t in txt if "manuels.html" in t)
        fin = sum(1 for t in txt if "veritas-school.com" in t and "manuels.html" not in t)
        W("| Bord %s | %s | %d |" % (niv, "oui" if pref else "**non — à ajouter**", fin))
    for f, lab in (("Bord_2nde_A_CORRIGE.docx", "Bord 2ⁿᵈᵉ"), ("Bord_1ere_A_CORRIGE.docx", "Bord 1ʳᵉ"),
                   ("Bord_Tle_A_CORRIGE.docx", "Bord Tˡᵉ")):
        d = Document(os.path.join(FIN2, f))
        txt = [p.text for p in d.paragraphs]
        pref = sum(1 for t in txt if "manuels.html" in t)
        fin = sum(1 for t in txt if "veritas-school.com" in t and "manuels.html" not in t)
        W("| %s | %s | %d |" % (lab, "oui" if pref else "**non — à ajouter**", fin))
    W("")

    # Pages publiées
    W("## 6. Pages publiées (`corriges/`)")
    W("")
    pages = sorted(glob.glob(os.path.join(ROOT, "corriges", "**", "*.html"), recursive=True))
    W("- **%d pages HTML** + `sitemap-corriges.xml` + `assets/corriges.css`" % len(pages))
    W("- Générées par `tools/build_corriges.py` (à relancer après toute modification des Guides ou des packs)")
    W("- Déploiement : `corriges/**` déclaré dans `.github/workflows/deploy.yml` (trigger + copie)")
    W("- Référencement : `corriges/sitemap-corriges.xml` ajouté à `sitemap-index.xml` et à `robots.txt`")
    W("")

    out = os.path.join(ROOT, "AUDIT_CORRIGES.md")
    with open(out, "w", encoding="utf-8") as f:
        f.write("\n".join(L) + "\n")
    print("✓ écrit :", out, "(%d lignes)" % len(L))

if __name__ == "__main__":
    main()
