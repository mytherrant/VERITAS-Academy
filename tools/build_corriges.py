# -*- coding: utf-8 -*-
"""
build_corriges.py — génère les pages publiques de corrigés des manuels VÉRITAS.

Sortie : corriges/  (hub + 1 index par niveau + 1 page par séquence + sitemap)

Sources de vérité
-----------------
1er cycle (6e, 5e, 4e, 3e) : Guide pédagogique .docx (styles T_Seq / T_Sem / T_Disc /
    Rubrique / Exercice / Corrige). Le docx est le master édité à la main.
2nd cycle (2nde, 1ere, Tle) : sources balisées pack/content/*.md (#SEQ #SEM #DISC
    #RUB #Q #EXO #SOL #SOLM). La numérotation reproduit EXACTEMENT celle de
    render_pack.py (q_n remis à 0 sur #RUB/#DISC/#SEM/#SEQ, exo_n sur #DISC/#SEM/#SEQ,
    « Leçon N · » sauf épreuve/bilan/évaluation/intégration/banque/diagnostic).

Ce qui est publié : énoncé de l'exercice + corrigé. Ce qui ne l'est PAS : les textes
d'auteur, les leçons (« Je retiens »), les définitions, les notes « Pour l'enseignant »
— ils restent la valeur du cahier imprimé et du Guide.

Usage : python tools/build_corriges.py
"""
import os, re, sys, glob, io, html, json, datetime
from collections import Counter

if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

try:
    from docx import Document
except ImportError:
    print("!! python-docx requis : pip install python-docx"); sys.exit(1)

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
OUT = os.path.join(ROOT, "corriges")
DESK = r"C:\Users\Mythe Errant\Desktop\Manuels"
FIN1 = os.path.join(DESK, "FINAUX_1er_cycle")
SITE = "https://veritas-school.com"
TODAY = datetime.date.today().isoformat()

# ─────────────────────────────────────────────────────────────── métadonnées
NIVEAUX = [
    dict(slug="6e",   label="6ᵉ",  long="Classe de 6ᵉ",        cycle="1er", examen=None,
         manuel="Mon Cahier de français 6ᵉ", livret="Livret d'activités 6ᵉ", kw="6eme sixieme"),
    dict(slug="5e",   label="5ᵉ",  long="Classe de 5ᵉ",        cycle="1er", examen=None,
         manuel="Mon Cahier de français 5ᵉ", livret="Livret d'activités 5ᵉ", kw="5eme cinquieme"),
    dict(slug="4e",   label="4ᵉ",  long="Classe de 4ᵉ",        cycle="1er", examen=None,
         manuel="Mon Cahier de français 4ᵉ", livret="Livret d'activités 4ᵉ", kw="4eme quatrieme"),
    dict(slug="3e",   label="3ᵉ",  long="Classe de 3ᵉ",        cycle="1er", examen="BEPC",
         manuel="Mon Cahier de français 3ᵉ", livret="Livret d'activités 3ᵉ", kw="3eme troisieme BEPC"),
    dict(slug="2nde", label="2ⁿᵈᵉ", long="Classe de 2ⁿᵈᵉ A",   cycle="2nd", examen=None,
         manuel="Mon Cahier de français 2ⁿᵈᵉ A",
         livret="Livret d'activités 2ⁿᵈᵉ A", kw="seconde 2nde A"),
    dict(slug="1ere", label="1ʳᵉ", long="Classe de 1ʳᵉ A",     cycle="2nd", examen="Probatoire",
         manuel="Mon Cahier de français 1ʳᵉ A",
         livret="Livret d'activités 1ʳᵉ A", kw="premiere 1ere A probatoire"),
    dict(slug="tle",  label="Tˡᵉ", long="Terminale A",         cycle="2nd", examen="BAC",
         manuel="Mon Cahier de français Terminale A",
         livret="Livret d'activités Terminale A", kw="terminale Tle A baccalaureat"),
]
SRC2ND = {"2nde": "Manuel_2nde_A", "1ere": "Manuel_1ere_A", "tle": "Manuel_Tle_A"}
# page « niveaux/ » correspondante (maillage interne, pas de doublon d'intention)
NIVEAU_SEO = {"6e": "francais-6eme", "5e": "francais-5eme", "4e": "francais-4eme",
              "3e": "francais-3eme", "2nde": "francais-seconde", "1ere": "francais-premiere",
              "tle": "francais-terminale"}

# ─────────────────────────────────────────────────────────────── utilitaires
INLINE = re.compile(r"(\*[^*]+\*|__[^_]+__|\|[^|]+\|)")

def esc(s):
    return html.escape(s or "", quote=False)

def iconifier(html):
    """Aucun émoji ne survit dans la page : les uns deviennent des icônes, les autres partent.

    Appelée en sortie d'`inline()`, donc sur du texte déjà échappé : les balises que
    nous venons d'écrire ne contiennent aucun émoji, la substitution ne peut pas les
    abîmer. Les signes typographiques (flèches, chiffres cerclés, ≈, ≠) ne sont pas
    des émojis et ne sont jamais touchés — ils portent du sens dans les schémas.
    """
    html = RE_ICONE.sub(lambda m: ico(ICONES[m.group(1)]), html)
    html = RE_EMOJI_NU.sub("", html)
    return re.sub(r"\s{2,}", " ", html).strip()


def inline(s):
    """Reproduit add_runs() de render_pack.py : *terme* = mis en avant, __x__ = italique."""
    out = []
    for chunk in INLINE.split(s or ""):
        if not chunk:
            continue
        if chunk.startswith("*") and chunk.endswith("*") and len(chunk) > 1:
            out.append('<b class="k">%s</b>' % esc(chunk[1:-1]))
        elif chunk.startswith("__") and chunk.endswith("__") and len(chunk) > 3:
            out.append("<em>%s</em>" % esc(chunk[2:-2]))
        elif chunk.startswith("|") and chunk.endswith("|") and len(chunk) > 1:
            out.append('<span class="hw">%s</span>' % esc(chunk[1:-1]))
        else:
            out.append(esc(chunk))
    return iconifier("".join(out))

def strip_prefix(t, *prefixes):
    for p in prefixes:
        if t.startswith(p):
            return t[len(p):].lstrip(" —:·-").strip()
    return t

TITRES_1ER = {
    "methodes": "Démarches méthodologiques officielles (MINESEC — APC-ESV)",
    "vers-le-bepc": "Préparation au BEPC — méthode, sujets zéro et épreuves blanches",
    "vers-le-cap": "Préparation au CAP — épreuves blanches",
}

def page_slug_1er(titre):
    m = re.match(r"S[ÉE]QUENCE\s+(\d+)", titre.strip(), re.I)
    if m:
        return "sequence-%s" % m.group(1)
    up = titre.upper()
    if "DÉMARCHE" in up or "DEMARCHE" in up:
        return "methodes"
    if "BEPC" in up:
        return "vers-le-bepc"
    if "CAP" in up:
        return "vers-le-cap"
    return re.sub(r"[^a-z0-9]+", "-", titre.lower()).strip("-")[:40] or "annexe"

def titre_propre(t):
    """« SÉQUENCE 3 — Citoyenneté… » → « Séquence 3 — Citoyenneté… » (les guides sont en capitales)."""
    t = t.strip()
    m = re.match(r"^S[ÉE]QUENCE\s+(\d+)\s*(.*)$", t)
    if m:
        reste = m.group(2).lstrip(" —–-·:").strip()
        return "Séquence %s — %s" % (m.group(1), reste) if reste else "Séquence %s" % m.group(1)
    if t.isupper() or (sum(c.isupper() for c in t if c.isalpha()) > 0.8 * max(1, sum(c.isalpha() for c in t))):
        # « BANQUE DE SUJETS… » → « Banque de sujets… »
        return t[:1] + t[1:].lower()
    return t

def num(n):
    return format(n, ',').replace(',', ' ')

MOTS_OUTILS = {"de", "du", "des", "d'", "la", "le", "les", "l'", "un", "une", "et", "ou",
               "à", "au", "aux", "en", "dans", "sur", "pour", "par", "avec", "son", "sa", "ses"}

def court(t, n=58):
    """Tronque proprement au mot (pour la balise <title>).

    Une troncature qui s'arrête sur « de » ou « la » se lit mal : on retire donc les
    mots outils restés en fin de coupe avant de poser les points de suspension.
    """
    t = t.strip()
    if len(t) <= n:
        return t
    mots = t[:n].rsplit(" ", 1)[0].split(" ")
    while len(mots) > 1 and mots[-1].lower().strip(",;:—–-") in MOTS_OUTILS:
        mots.pop()
    cut = " ".join(mots).rstrip(" ,;:—–-")
    return cut + "…"

def libelle_competence(comp):
    """« À la fin de la séquence, l'élève produira un résumé… » → « Produire un résumé… »"""
    c = re.sub(r"^À la fin de la séquence,?\s*", "", comp or "").strip()
    c = re.sub(r"^l['’](élève|apprenant)\s+", "", c)
    c = re.sub(r"^produira\b", "Produire", c)
    c = re.sub(r"^utilisera\b", "Utiliser", c)
    c = c.rstrip(".")
    return (c[:1].upper() + c[1:]) if c else ""


# ── Icônes vectorielles (sprite /assets/veritas-icons.svg) ────────────────────
# Les cahiers imprimés balisent leurs rubriques par des émojis (🔍 Je repère,
# ✍️ Exercice, 🎯 Objectif…). À l'écran, on les rend par de VRAIES icônes
# vectorielles : même repère visuel, rendu net et homogène sur tous les appareils.
# Règle : on ne convertit QUE l'émoji de tête (marqueur de rubrique). Un émoji au
# milieu d'une phrase appartient au texte de l'exercice — il n'est jamais touché.
ICONES = {
    "🎯": "i-target", "✍": "i-pen", "🔬": "i-flask", "📊": "i-chart", "🔍": "i-search",
    "💡": "i-lightbulb", "📖": "i-book-open", "📝": "i-pen", "✅": "i-check", "🖋": "i-pen",
    "✏": "i-pen", "🔄": "i-refresh", "🪶": "i-feather", "❌": "i-x-circle", "🎙": "i-mic",
    "🧭": "i-compass", "⚖": "i-scale", "🔗": "i-link", "🚀": "i-arrow-right", "💬": "i-message",
    "🔤": "i-type", "🔠": "i-type", "🔡": "i-type", "🎒": "i-backpack", "🧮": "i-calculator", "📘": "i-book", "🔎": "i-search",
    "👑": "i-award", "🔊": "i-volume", "🔧": "i-tool", "🛠": "i-tool", "🔢": "i-grid",
    "🔁": "i-refresh", "🎲": "i-dice", "👀": "i-eye", "🪜": "i-steps", "🧩": "i-puzzle",
    "📕": "i-book", "🎤": "i-mic", "✂": "i-scissors", "📚": "i-book-open", "📢": "i-megaphone",
    "📣": "i-megaphone", "🎭": "i-users", "📲": "i-scan", "📌": "i-quote", "🌐": "i-globe",
    "🗣": "i-mic", "🧑": "i-teacher", "🏫": "i-school", "⚠": "i-warning", "🎨": "i-image",
    "⚡": "i-sparkle", "🎵": "i-music", "🎼": "i-music", "📈": "i-chart", "🖼": "i-image",
    "🧱": "i-grid", "🌀": "i-refresh", "🧐": "i-eye", "🧾": "i-file-text", "🌳": "i-map",
    "🖊": "i-pen", "📓": "i-notebook", "💻": "i-graduation", "🤖": "i-bot", "🎬": "i-play",
    "🧵": "i-link", "🌍": "i-globe", "🎢": "i-sparkle", "🕐": "i-clock", "🩺": "i-gauge",
    "💪": "i-award", "➕": "i-grid", "❓": "i-lightbulb", "✒": "i-pen", "🖇": "i-link",
    "🌡": "i-gauge", "🔀": "i-refresh", "🎮": "i-gamepad", "🔐": "i-lock", "🔑": "i-key",
    "🗓": "i-calendar", "🔭": "i-eye", "📊": "i-chart", "👩": "i-users", "📗": "i-book", "📒": "i-book", "📐": "i-grid",
    "🌺": "i-quote", "🧠": "i-brain", "👉": "i-arrow-right", "🎓": "i-graduation",
    "🌟": "i-star", "📜": "i-file-text", "🤓": "i-graduation",
    "🔹": "i-badge", "☐": "i-box", "⏳": "i-clock", "📄": "i-file-text",
}
EMOJI = re.compile("[🌀-🫿☀-➿⬀-⯿]")
# Aucun émoji ne doit atteindre la page : ceux de la table deviennent des icônes
# vectorielles, les autres sont retirés. On ne touche PAS aux signes typographiques
# (→, ①②③, ≈, ≠, ↓) : ce sont des symboles de contenu, pas des marqueurs.
RE_ICONE = re.compile("(" + "|".join(sorted((re.escape(k) for k in ICONES), key=len, reverse=True))
                      + ")️?")
RE_EMOJI_NU = re.compile("[🌀-🫿⬀-⯿]️?")


def ico(name, cls="i"):
    return ('<svg class="%s" aria-hidden="true" focusable="false">'
            '<use href="/assets/veritas-icons.svg#%s"/></svg>' % (cls, name))

def pull_icon(t, defaut=None, fenetre=26):
    """Extrait l'émoji-marqueur de tête → (html d'icône, texte nettoyé).

    `fenetre` : on ne cherche que dans les premiers caractères, pour ne jamais
    toucher à un émoji qui appartient au contenu (leçons sur le langage SMS…).
    """
    t = t or ""
    m = EMOJI.search(t[:fenetre])
    if not m:
        return (ico(defaut) if defaut else ""), t
    nom = ICONES.get(m.group(0))
    if not nom:
        return (ico(defaut) if defaut else ""), t
    reste = (t[:m.start()] + t[m.end():]).replace("️", "")
    reste = re.sub(r"\s{2,}", " ", reste).strip(" ·—-").strip()
    # « Leçon 1 ·  Grammaire » → « Leçon 1 · Grammaire »
    reste = reste.replace(" · ·", " ·").replace("·  ", "· ")
    return ico(nom), reste

class Page:
    def __init__(self, slug, titre):
        self.slug, self.titre, self.lecons = slug, titre, []
        self.comp = ""
        self.rang = 0
    @property
    def n_items(self):
        return sum(len(l["items"]) for l in self.lecons)

class Lecon:
    @staticmethod
    def new(titre, objectif=None, semaine=None):
        return dict(titre=titre, objectif=objectif, semaine=semaine, items=[])

# ───────────────────────────────────────────────── extraction 1er cycle (docx)
def extract_1er(niv):
    path = os.path.join(FIN1, "%s_Guide_pedagogique.docx" % niv["slug"])
    doc = Document(path)
    pages, by_slug = [], {}
    page = lec = rub = item = None
    semaine = None
    stats = Counter()
    for p in doc.paragraphs:
        st = p.style.name if p.style is not None else ""
        t = p.text.strip()
        if not t:
            continue
        if st == "T_Seq":
            slug = page_slug_1er(t)
            t = TITRES_1ER.get(slug, titre_propre(t))
            if slug in by_slug:
                page = by_slug[slug]
            else:
                page = Page(slug, t); by_slug[slug] = page; pages.append(page)
            lec = rub = item = None; semaine = None
        elif st == "T_Sem":
            semaine = strip_prefix(t, "👉")
        elif st == "T_Disc":
            if page is None:
                continue
            lec = Lecon.new(t, semaine=semaine)
            page.lecons.append(lec); rub = item = None
        elif st == "Objectif" and lec is not None and not lec["objectif"]:
            lec["objectif"] = strip_prefix(t, "🎯")
        elif st == "Rubrique":
            rub = t; item = None
        elif st == "Exercice":
            if lec is None:
                continue
            item = dict(rub=rub, q=strip_prefix(t, "✍️").strip(), sub=[], sol=[])
            lec["items"].append(item); stats["exos"] += 1
        elif st in ("Liste", "Liste1") and item is not None and not item["sol"]:
            item["sub"].append(strip_prefix(t, "•"))
        elif st == "Corrige" and item is not None:
            item["sol"].append(strip_prefix(t, "✅ Corrigé —", "✅ Corrigé", "✅"))
            stats["corriges"] += 1
    # nettoyage : on ne publie que ce qui a un corrigé
    for pg in pages:
        for l in pg.lecons:
            l["items"] = [i for i in l["items"] if i["sol"]]
        pg.lecons = [l for l in pg.lecons if l["items"]]
    pages = [pg for pg in pages if pg.lecons]
    return pages, stats

# ───────────────────────────────────────────────── extraction 2nd cycle (md)
TAG = re.compile(r"^#([A-Z_0-9]+)::\s?(.*)$")
NO_LECON = re.compile(r"épreuve|remédiation|bilan|évaluation|intégration|banque de sujets|sujet d'entra", re.I)

def extract_2nd(niv):
    cdir = os.path.join(DESK, SRC2ND[niv["slug"]], "pack", "content")
    files = sorted(glob.glob(os.path.join(cdir, "*.md")))
    pages, by_slug = [], {}
    stats = Counter()
    page = lec = rub = item = None
    semaine = None
    lecon_n = q_n = exo_n = 0
    diag = False

    def get_page(slug, titre):
        nonlocal page
        if slug in by_slug:
            page = by_slug[slug]
        else:
            page = Page(slug, titre); by_slug[slug] = page; pages.append(page)
        return page

    for f in files:
        for raw in open(f, encoding="utf-8"):
            m = TAG.match(raw.rstrip("\n"))
            if not m:
                continue
            tag, val = m.group(1), m.group(2).strip()
            if tag == "PART":
                low = val.lower()
                if "diagnost" in low:
                    get_page("diagnostic", "Évaluation diagnostique")
                elif "sujet" in low or "banque" in low:
                    get_page("sujets-examen", "Sujets d'examen — banque corrigée et entraînement")
                else:
                    get_page(re.sub(r"[^a-z0-9]+", "-", low).strip("-")[:32] or "annexe", val)
                lec = rub = item = None
                lecon_n = q_n = exo_n = 0
            elif tag == "SEQ":
                num = val.split("::", 1)[0].strip()
                comp = val.split("::", 1)[1].strip() if "::" in val else ""
                lib = libelle_competence(comp)
                get_page("sequence-%s" % num, "Séquence %s — %s" % (num, lib) if lib else "Séquence %s" % num)
                page.comp = comp
                lec = rub = item = None; semaine = None
                lecon_n = q_n = exo_n = 0; diag = False
            elif tag == "SEM":
                parts = [x.strip() for x in val.split("::", 1)]
                semaine = ("Semaine %s — %s" % (parts[0], parts[1])) if len(parts) > 1 else ("Semaine %s" % parts[0])
                lecon_n = q_n = exo_n = 0; diag = False
            elif tag == "DISC":
                if page is None:
                    get_page("diagnostic", "Évaluation diagnostique")
                q_n = exo_n = 0
                label = val
                if re.search(r"diagnost", val, re.I):
                    diag = True
                if not diag and not NO_LECON.search(val):
                    lecon_n += 1
                    label = "Leçon %d · %s" % (lecon_n, val)
                lec = Lecon.new(label, semaine=semaine)
                page.lecons.append(lec); rub = item = None
            elif tag == "OBJ" and lec is not None and not lec["objectif"]:
                lec["objectif"] = val
            elif tag == "RUB":
                rub = val; q_n = 0; item = None
            elif tag in ("Q", "QW", "QN"):
                if lec is None:
                    continue
                q_n += 1
                # Les épreuves d'examen numérotent déjà leurs questions dans la source
                # (« #Q:: 1. a) … ») : ne pas préfixer une seconde fois, sinon « 1. 1. a) ».
                deja_num = re.match(r"^\d+\s*[.)]\s", val) is not None
                item = dict(rub=rub, q=(val if deja_num else "%d. %s" % (q_n, val)), sub=[], sol=[])
                lec["items"].append(item); stats["questions"] += 1
            elif tag == "EXO":
                if lec is None:
                    continue
                parts = [x.strip() for x in val.split("::", 1)]
                cons = parts[1] if len(parts) > 1 else parts[0]
                exo_n += 1
                item = dict(rub=rub, q="Exercice %d : %s" % (exo_n, cons), sub=[], sol=[])
                lec["items"].append(item); stats["exos"] += 1
            elif tag == "LI" and item is not None and not item["sol"]:
                item["sub"].append(val)
            elif tag in ("SOL", "SOLM"):
                if item is None:
                    continue
                item["sol"].append(val); stats["corriges"] += 1
    for pg in pages:
        for l in pg.lecons:
            l["items"] = [i for i in l["items"] if i["sol"]]
        pg.lecons = [l for l in pg.lecons if l["items"]]
    pages = [pg for pg in pages if pg.lecons]
    # ordre : diagnostic, séquences, annexes
    def order(pg):
        if pg.slug == "diagnostic":
            return (0, 0)
        m = re.match(r"sequence-(\d+)", pg.slug)
        return (1, int(m.group(1))) if m else (2, 0)
    pages.sort(key=order)
    return pages, stats

# ────────────────────────────────────── extraction 1er cycle — CAHIER (« Bord »)
# Deux ouvrages distincts par niveau, donc deux axes de corrigés :
#   • le LIVRET d'activités  → corrigés dans le Guide, publiés en `sequence-N.html` ;
#   • le CAHIER de français  → le « Bord », publié ici en `cahier-module-N.html`.
# Recouvrement mesuré entre les deux : 0 %. Les corrigés du cahier n'existent nulle
# part ailleurs : ils sont rédigés à la main dans content/corriges-cahier/{niv}/,
# l'ossature venant de _bord_extract/{niv}.json (cf. tools/extract_bord.py).
EXTRAIT = os.path.join(ROOT, "_bord_extract")
SRC_CAHIER = os.path.join(ROOT, "content", "corriges-cahier")
ROMAINS = {"I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6, "VII": 7, "VIII": 8}
RE_MODULE = re.compile(r"MODULE\s+([IVX]+|\d+)\s*[:\-—–]?\s*(.*)$", re.I)
# Le 1er cycle découpe l'année en modules, le 2nd cycle en séquences didactiques :
# même ossature, deux noms. Le slug suit le mot que l'élève lit sur son cahier.
RE_SEQ = re.compile(r"S[ÉE]QUENCE\s+(\d+)\s*[:\-—–]?\s*(.*)$", re.I)
EMOJI_TETE = re.compile(r"^[\W_]*?([🌀-🫿☀-➿⬀-⯿])\s*")


def module_page(titre):
    """« MODULE III — CITOYENNETÉ… » → (slug, « Module 3 — Citoyenneté… », rang)."""
    s = RE_SEQ.search(titre or "")
    if s:
        n = int(s.group(1))
        reste = s.group(2).strip(" :—–-")
        return ("cahier-sequence-%d" % n,
                ("Séquence %d — %s" % (n, reste)) if reste else "Séquence %d" % n, n)
    m = RE_MODULE.search(titre or "")
    if not m:
        return "cahier-evaluation-diagnostique", "Évaluation diagnostique de rentrée", 0
    n = m.group(1).upper()
    n = ROMAINS.get(n, None) or int(n) if not n.isdigit() else int(n)
    reste = m.group(2).strip(" :—–-")
    reste = reste[:1].upper() + reste[1:].lower() if reste.isupper() else reste
    return ("cahier-module-%d" % n,
            ("Module %d — %s" % (n, reste)) if reste else "Module %d" % n, n)


def repere_cahier(txt, kind, num=None):
    """Le repère que l'élève retrouve dans son cahier : « Exercice 2 », « Question 5 »…

    Les cahiers du 2nd cycle numérotent leurs questions dans la mise en page, pas dans
    l'énoncé : l'extracteur transporte ce numéro (`num`), et il fait foi — c'est celui
    que l'élève a sous les yeux.
    """
    if num:
        n = str(num).strip()
        return n if re.match(r"^(Exercice|Épreuve|Epreuve|Sujet)\b", n, re.I) else "Question %s" % n
    t = EMOJI_TETE.sub("", (txt or "").strip()).lstrip("️ ")
    m = re.match(r"^(?:✍️\s*)?(Exercice|Exo)\s*(\d+)", t, re.I)
    if m:
        return "Exercice %s" % m.group(2)
    m = re.match(r"^Sujet\s*(\d+)", t, re.I)
    if m:
        return "Sujet %s" % m.group(1)
    m = re.match(r"^(\d{1,2})\s*[°.)\-–]", t)
    if m:
        return "Question %s" % m.group(1)
    if kind == "prod":
        return "Production"
    return "Question"


def lire_corriges(niv):
    """content/corriges-cahier/{niv}/module-N.md → ({item_id: [paragraphes]}, {item_id: libellé}).

    Deux balises. `#SOL:: id :: texte` attache un corrigé à un item du cahier ; répétée,
    elle ajoute un paragraphe. `#LAB:: id :: libellé` nomme un item que l'extracteur ne
    pouvait pas voir — une épreuve dont la consigne tient dans le texte, sans énoncé
    numéroté. Ces items-là portent un identifiant en `-x1`, `-x2`… et se rangent à la
    fin de leur leçon.
    """
    sols, labs = {}, {}
    d = os.path.join(SRC_CAHIER, niv)
    if not os.path.isdir(d):
        return sols, labs
    for f in sorted(glob.glob(os.path.join(d, "module-*.md"))):
        for raw in open(f, encoding="utf-8"):
            if not (raw.startswith("#SOL::") or raw.startswith("#LAB::")):
                continue
            parts = raw.rstrip("\n").split("::", 2)
            if len(parts) < 3:
                continue
            if raw.startswith("#LAB::"):
                labs[parts[1].strip()] = parts[2].strip()
            else:
                sols.setdefault(parts[1].strip(), []).append(parts[2].strip())
    return sols, labs


def extract_cahier(niv):
    """Pages du cahier : une par module, alimentées par les corrigés rédigés."""
    path = os.path.join(EXTRAIT, "%s.json" % niv["slug"])
    if not os.path.exists(path):
        return [], Counter()
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    sols, labs = lire_corriges(niv["slug"])
    # Modules dont les corrigés ont été rédigés : pour ceux-là, on ignore les blocs
    # « ✅ Corrigé » que le cahier porte déjà (évaluations diagnostiques), sinon la page
    # afficherait deux fois la même correction — une fois item par item, une fois en bloc.
    rediges = set()
    for k in sols:
        m = re.match(r"[^-]+-m(\d+)-", k)
        if m:
            rediges.add(int(m.group(1)))
    stats = Counter()
    pages, by_slug = [], {}
    for mod in data["modules"]:
        slug, titre, rang = module_page(mod["titre"])
        if slug in by_slug:
            page = by_slug[slug]
        else:
            page = Page(slug, titre); page.rang = rang
            by_slug[slug] = page; pages.append(page)
        for l in mod["lecons"]:
            lec = Lecon.new(l["titre"], objectif=strip_prefix(l["objectif"] or "", "🎯"),
                            semaine=l["semaine"])
            for it in l["items"]:
                integre = [x for x in it.get("sol", []) if x]
                if mod["n"] in rediges:
                    integre = []
                sol = sols.get(it["id"]) or integre
                if not sol:
                    continue          # pas de corrigé rédigé → pas publié
                lec["items"].append(dict(rub=it["rub"], q=it["txt"], sub=[], sol=sol,
                                         lab=repere_cahier(it["txt"], it["kind"], it.get("num"))))
                stats["corriges"] += 1
            # items ajoutés à la main (épreuves sans énoncé numéroté) : « …-l3-x1 »
            prefixe = "%s-m%d-l%d-x" % (niv["slug"], mod["n"], l["n"])
            for k in sorted(x for x in sols if x.startswith(prefixe)):
                lec["items"].append(dict(rub="", q="", sub=[], sol=sols[k],
                                         lab=labs.get(k, "Exercice")))
                stats["corriges"] += 1
            if lec["items"]:
                page.lecons.append(lec)
    pages = [p for p in pages if p.lecons]
    pages.sort(key=lambda p: p.rang)
    return pages, stats


# ─────────────────────────────────────────────────────────────── gabarit HTML
# CSS et JS : /assets/veritas-pages.css et /assets/veritas-ui.js, communs au site.
#
# ⚠️ DEUX PIÈGES, ET ILS SE PAIENT LES DEUX EN SILENCE.
#
# 1. LE CACHE-BUSTER N'EST PAS FACULTATIF. .htaccess sert les CSS avec
#    « Cache-Control: public, max-age=31536000, immutable ». Une feuille
#    demandée sans « ?v= » est donc gardée UN AN par le navigateur : la
#    correction déployée aujourd'hui n'atteindra jamais un visiteur déjà venu,
#    et rien ne le signalera — la page s'affiche, simplement avec l'ancienne
#    charte. Le numéro est lu dans la coquille pour rester aligné sur le reste
#    du site à chaque déploiement.
#
# 2. LA POLICE DOIT ÊTRE DEMANDÉE. veritas-pages.css déclare Poppins depuis
#    l'unification de la charte, mais une feuille ne télécharge pas une police
#    Google : il faut le <link>. Sans lui, ces 56 pages retombaient sur le
#    repli système (Segoe UI sur Windows) pendant que la vitrine et les 64
#    autres pages statiques rendaient en Poppins. C'est exactement l'écart
#    « la moitié du site sur l'ancienne charte » qu'on venait de corriger.
#    Chargement non bloquant (media=print puis onload) : ces pages sont des
#    entrées depuis Google, elles doivent peindre tout de suite.

def _version_assets():
    """Reprend la version demandée par la coquille (app.js?v=…)."""
    import re as _re
    coquille = os.path.join(ROOT, "VERITAS_v1.2.html")
    try:
        with open(coquille, "r", encoding="utf-8", errors="ignore") as f:
            tete = f.read(200000)
        m = _re.search(r"app\.js\?v=([0-9.]+)", tete)
        if m:
            return m.group(1)
    except OSError:
        pass
    # Pas de repli silencieux : une feuille sans version se fige un an en cache.
    raise SystemExit("build_corriges : version introuvable dans VERITAS_v1.2.html "
                     "(motif app.js?v=…). Corriger avant de régénérer.")

VER = _version_assets()

def head(title, desc, canonical, depth=1, jsonld=""):
    # depth = profondeur SOUS corriges/ (0 = hub, 1 = pages de niveau)
    up = "../" * max(0, depth - 1)
    return """<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>%(title)s</title>
<meta name="description" content="%(desc)s">
<meta name="robots" content="index,follow">
<meta property="og:type" content="article">
<meta property="og:site_name" content="Centre VÉRITAS">
<meta property="og:title" content="%(title)s">
<meta property="og:description" content="%(desc)s">
<meta property="og:url" content="%(can)s">
<meta property="og:image" content="%(site)s/uploads/logo-veritas.png">
<link rel="canonical" href="%(can)s">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Poppins:wght@400;500;600&display=swap" media="print" onload="this.media='all'">
<noscript><link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Poppins:wght@400;500;600&display=swap"></noscript>
<link rel="stylesheet" href="/assets/veritas-pages.css?v=%(ver)s">
%(jsonld)s
</head>
<body>
""" % dict(title=esc(title), desc=esc(desc), can=canonical, site=SITE, jsonld=jsonld, ver=VER)

# ───────────────────────────────────────────────────── barre de recherche
# UNE page de séquence porte jusqu'à 54 corrigés répartis en sept leçons.
# L'élève arrive avec son cahier ouvert sur « exercice 12 » ou « le complément
# d'objet » : il faisait défiler. La barre filtre en direct les leçons ET les
# corrigés, ouvre ce qui correspond, et surligne le terme trouvé.
#
# Pourquoi le script est EN LIGNE et non dans /assets : ces pages ne chargent
# pas toutes le même JS (109 sur 128 ont veritas-convert.js, 67 veritas-ui.js,
# 18 aucun). Une recherche qui ne marcherait que sur une partie des corrigés
# serait pire que pas de recherche — l'élève conclurait que le corrigé n'existe
# pas. Une centaine de lignes, aucune requête, aucune dépendance.
#
# La normalisation retire les accents : « élève » se tape « eleve » sur un
# clavier de téléphone, et « numero 12 » doit trouver « n° 12 ».
SEARCHBAR = (
    '<div class="cfind">'
    '<label class="cfind-l" for="cq">Trouver un corrigé ou une leçon</label>'
    '<div class="cfind-box">'
    '<svg class="i cfind-i" aria-hidden="true" focusable="false">'
    '<use href="/assets/veritas-icons.svg#i-search"></use></svg>'
    '<input id="cq" class="cfind-in" type="search" autocomplete="off" '
    'placeholder="Numéro d\'exercice, mot de la leçon… (ex. « 12 », « accord », « COD »)" '
    'aria-describedby="cfind-say">'
    '<button type="button" class="cfind-x" hidden aria-label="Effacer la recherche">' + ico("i-x-circle") + '</button>'
    '</div>'
    '<p class="cfind-say" id="cfind-say" role="status" aria-live="polite"></p>'
    '</div>'
    # ⚠️ Attendre le DOM. Ce <script> est placé juste sous la barre, donc AVANT
    # les <section class="lec"> qu'il doit indexer : exécuté tel quel, il
    # construisait un index vide et la recherche ne trouvait jamais rien — sans
    # la moindre erreur en console, ce qui est le pire des deux mondes.
    '<script>(function(){function go(){'
    'var q=document.getElementById("cq");if(!q)return;'
    'var x=document.querySelector(".cfind-x"),say=document.getElementById("cfind-say");'
    'var secs=[].slice.call(document.querySelectorAll("section.lec"));'
    'if(!secs.length)return;'
    'var nSol=document.querySelectorAll("details.sol").length;'
    'function norm(s){return (s||"").toLowerCase()'
    '.normalize("NFD").replace(/[\\u0300-\\u036f]/g,"")'
    '.replace(/[\\u2019\\u00b0]/g," ").replace(/\\s+/g," ").trim();}'
    # Index construit UNE fois : refaire textContent sur 54 blocs à chaque
    # frappe rame visiblement sur un téléphone d'entrée de gamme.
    'var idx=secs.map(function(s){'
    'var sols=[].slice.call(s.querySelectorAll("details.sol"));'
    'return {sec:s,t:norm(s.textContent),'
    'sols:sols.map(function(d){var p=d.previousElementSibling;'
    'return {el:d,t:norm((p?p.textContent:"")+" "+d.textContent),prev:p};})};});'
    'function filtre(){'
    'var v=norm(q.value);x.hidden=!q.value;'
    'if(!v){idx.forEach(function(o){o.sec.hidden=false;'
    'o.sols.forEach(function(s){s.el.hidden=false;if(s.prev)s.prev.hidden=false;s.el.open=false;});});'
    'say.textContent="";return;}'
    'var mots=v.split(" ").filter(Boolean),n=0,nl=0,nt=0;'
    'idx.forEach(function(o){'
    'var vus=0;'
    'o.sols.forEach(function(s){'
    'var ok=mots.every(function(m){return s.t.indexOf(m)>=0;});'
    's.el.hidden=!ok;if(s.prev)s.prev.hidden=!ok;s.el.open=ok;if(ok){vus++;n++;}});'
    # Une leçon dont le TITRE correspond reste entière : on cherche parfois la
    # leçon, pas un exercice précis.
    'var titre=mots.every(function(m){return o.t.indexOf(m)>=0;});'
    'if(!vus&&titre){o.sols.forEach(function(s){s.el.hidden=false;if(s.prev)s.prev.hidden=false;});nt++;}'
    'var garde=vus>0||titre;o.sec.hidden=!garde;if(garde)nl++;});'
    # Le compte doit dire la vérité de ce qu'on VOIT : une leçon retenue sur son
    # titre reste entière, on affiche donc plus de corrigés qu'il n'y a de
    # correspondances. Annoncer « 1 corrigé » devant 45 blocs ouverts ferait
    # douter du filtre plutôt que du mot cherché.
    'say.textContent=(n||nt)?('
    '(n?n+" corrig"+(n>1?"\\u00e9s":"\\u00e9")+" trouv"+(n>1?"\\u00e9s":"\\u00e9"):"")'
    '+(n&&nt?" \\u00b7 ":"")'
    '+(nt?nt+" le\\u00e7on"+(nt>1?"s":"")+" affich\\u00e9e"+(nt>1?"s":"")+" en entier":"")'
    '):"Rien avec \\u00ab "+q.value+" \\u00bb. Essaie le num\\u00e9ro seul, ou un mot de la le\\u00e7on.";}'
    'q.addEventListener("input",filtre);'
    'q.addEventListener("keydown",function(e){if(e.key==="Escape"){q.value="";filtre();}});'
    'x.addEventListener("click",function(){q.value="";q.focus();filtre();});'
    # Arriver avec ?q=12 depuis le hub : la recherche est déjà faite.
    'var p=new URLSearchParams(location.search).get("q");'
    'if(p){q.value=p;filtre();q.scrollIntoView({block:"center"});}'
    '}'
    'if(document.readyState==="loading")document.addEventListener("DOMContentLoaded",go);else go();'
    '})();</script>'
)


def foot(depth=1):
    # veritas-convert.js ajoute le bloc de conversion AVANT ce pied de page.
    # Il doit rester ici : les pages de corrigés sont régénérées par ce script,
    # une balise ajoutée à la main dans le HTML serait effacée au prochain build.
    # Le pied de page mène désormais aussi aux abonnements : les deux seuls
    # liens sortants de /corriges/ pointaient tous les deux vers manuels.html.
    return """<footer class="bot">© <span id="y"></span> Centre VÉRITAS · Mythe Errant — Douala, Cameroun ·
<a href="%(site)s">veritas-school.com</a> · <a href="%(site)s/manuels.html">Espace Manuels</a> ·
<a href="%(site)s/#tarifs">Abonnements</a></footer>
<script src="/assets/veritas-ui.js?v=%(ver)s" defer></script>
<script src="/assets/veritas-convert.js?v=%(ver)s" defer></script>
<script src="/assets/ambassa.js?v=%(ver)s" defer></script>
</body>
</html>
""" % dict(site=SITE, ver=VER)

def jsonld_page(name, desc, url, level, crumbs):
    items = "".join(
        '{"@type":"ListItem","position":%d,"name":"%s","item":"%s"}%s' % (i + 1, esc(n), u, "," if i < len(crumbs) - 1 else "")
        for i, (n, u) in enumerate(crumbs))
    return ('<script type="application/ld+json">{"@context":"https://schema.org","@graph":['
            '{"@type":"LearningResource","name":"%s","description":"%s","url":"%s",'
            '"inLanguage":"fr","learningResourceType":"Corrigé d\'exercices",'
            '"educationalLevel":"%s","educationalUse":"self-study",'
            '"isAccessibleForFree":true,'
            '"provider":{"@type":"EducationalOrganization","name":"Centre VÉRITAS",'
            '"url":"%s","address":{"@type":"PostalAddress","addressLocality":"Douala","addressCountry":"CM"}}},'
            '{"@type":"BreadcrumbList","itemListElement":[%s]}]}</script>'
            % (esc(name), esc(desc), url, esc(level), SITE, items))

CTA = """
  <h2 class="sec">%(ar)sContinuer</h2>
  <div class="cta">
    <a href="/eleve/">%(sac)sMon espace élève<small>Toutes mes ressources au même endroit</small></a>
    <a href="%(site)s">%(bot)sProf. Ambassa (IA)<small>Poser une question sur un corrigé</small></a>
    <a href="/outils/">%(calc)sOutils gratuits<small>Moyenne, points manquants, planning</small></a>
    <a href="/constellation.html">%(comp)sLa Constellation VÉRITAS<small>Tout ce qui accompagne le cahier</small></a>
  </div>
""" % dict(site=SITE, ar=ico("i-arrow-right"), sac=ico("i-backpack"), bot=ico("i-bot"),
           calc=ico("i-calculator"), comp=ico("i-compass"))

# ─────────────────────────────────────────────────────────────── rendu pages
def repere(q):
    """Renvoie le SEUL repère de l'exercice : « Exercice 3 », « Question 2 »…

    Le cahier est vendu ; les corrigés sont offerts. On ne republie donc JAMAIS
    l'énoncé en ligne — l'élève le lit dans son cahier et vient chercher ici la
    correction correspondante. Ne sort de cette fonction qu'un numéro.
    """
    q = (q or "").strip()
    m = re.match(r"^(Exercice|Exo|Devoir|Tâche)\s*(\d+)", q, re.I)
    if m:
        return "Exercice %s" % m.group(2)
    m = re.match(r"^(\d+)\s*[.)]", q)
    if m:
        return "Question %s" % m.group(1)
    return "Question"

def render_items(lec):
    out = []
    last_rub = None
    n_sans_num = 0
    for it in lec["items"]:
        if it["rub"] and it["rub"] != last_rub:
            icn, rub = pull_icon(it["rub"], "i-quote")
            out.append('<p class="rub">%s<span>%s</span></p>' % (icn, inline(rub)))
            last_rub = it["rub"]
            n_sans_num = 0
        sol = ""
        for stxt in it["sol"]:
            # « 📊 Barème — … », « 💡 Astuce — … » en tête de paragraphe : même
            # traitement que les rubriques (icône), le reste du texte est intact.
            icn_s, corps = pull_icon(stxt)
            sol += "<p>%s%s</p>" % (icn_s, inline(corps))
        lab = it.get("lab") or repere(it["q"])
        if lab == "Question":                      # énoncé sans numéro : on en pose un
            n_sans_num += 1
            lab = "Question %d" % n_sans_num
        # L'énoncé (it["q"]) et ses options (it["sub"]) restent dans le cahier.
        out.append(
            '<div class="ex"><p class="q"><b>%s</b></p>'
            '<details class="sol"><summary>%sVoir le corrigé</summary>'
            '<div class="body">%s</div></details></div>'
            % (esc(lab), ico("i-check"), sol))
    return "\n".join(out)

def render_sequence_page(niv, page, pages, axe="livret"):
    n_ex = page.n_items
    cah = axe == "cahier"
    ouvrage = niv["manuel"] if cah else niv.get("livret", niv["manuel"])
    quoi = "cahier de français" if cah else "livret d'activités"
    ton = "ton cahier" if cah else "ton livret"
    exam = (" · préparation au %s" % niv["examen"]) if niv["examen"] else ""
    title = "Corrigés %s %s — %s | VÉRITAS" % (
        niv["label"], "Cahier" if cah else "Livret", court(page.titre, 34))
    desc = ("%s exercices corrigés du %s %s — %s. Programme MINESEC%s. "
            "Les énoncés restent dans %s."
            % (num(n_ex), "cahier" if cah else "livret", niv["label"], court(page.titre, 52).rstrip("…"),
               exam.rstrip(".").replace(" Préparation au", ", préparation au"), ton))
    url = "%s/corriges/%s/%s.html" % (SITE, niv["slug"], page.slug)
    crumbs = [("Accueil", SITE), ("Corrigés des manuels", SITE + "/corriges/"),
              (niv["long"], "%s/corriges/%s/" % (SITE, niv["slug"])), (page.titre, url)]
    h = head(title, desc, url, depth=2, jsonld=jsonld_page(title, desc, url, niv["long"], crumbs))
    nav = "".join(
        '<a href="%s.html"%s>%s</a>' % (p.slug, ' aria-current="page"' if p is page else "", esc(short_title(p.titre)))
        for p in pages)
    body = ['<header class="top"><span class="badge">Corrigés · %s · %s</span>'
            "<h1>%s — %s</h1><p>%s exercices corrigés du %s « %s ». "
            "Les énoncés restent dans %s ; ici, les corrections — et elles sont gratuites.</p></header>"
            % (esc(niv["long"]), "Cahier" if cah else "Livret", esc(niv["long"]), esc(page.titre),
               num(n_ex), quoi, esc(ouvrage), ton),
            '<div class="wrap">',
            '<div class="intro">' + ico("i-notebook", "i lg") + ' <strong>Comment t\'en servir ?</strong> Ouvre %s '
            "à la leçon, fais l'exercice, puis retrouve ici son numéro et clique sur « Voir le corrigé ». "
            "<strong>Les énoncés ne sont pas recopiés en ligne</strong> — ils sont dans %s. "
            "Ici, tu trouves les corrections, et elles sont gratuites.</div>" % (ton, ton),
            '<p class="crumb"><a href="%s/">Accueil</a> › <a href="../">Corrigés</a> › '
            '<a href="./">%s</a> › %s</p>' % (SITE, esc(niv["long"]), esc(page.titre)),
            '<nav class="seqnav">%s</nav>' % nav,
            SEARCHBAR,
            '<div class="tools"><button data-all="open">Tout ouvrir</button>'
            '<button data-all="close">Tout fermer</button><span class="sep"></span>'
            '<button data-a11y aria-pressed="false">' + ico("i-type") + 'Lecture facilitée</button>'
            '<button data-speak="page" aria-pressed="false">' + ico("i-volume")
            + 'Écouter les corrigés</button>'
            '<button data-share>' + ico("i-link") + 'Partager cette page</button></div>']
    cur_sem = object()
    for lec in page.lecons:
        body.append('<section class="lec">')
        if lec["semaine"] and lec["semaine"] != cur_sem:
            cur_sem = lec["semaine"]
            body.append('<p class="sem">%s</p>' % esc(lec["semaine"]))
        icn, titre = pull_icon(lec["titre"], "i-book-open")
        body.append("<h3>%s<span>%s</span></h3>" % (icn, inline(titre)))
        if lec["objectif"]:
            body.append('<p class="obj">%s<span>%s</span></p>'
                        % (ico("i-target"), inline(lec["objectif"])))
        body.append(render_items(lec))
        body.append("</section>")
    body.append(CTA)
    body.append('<p class="note">Ces corrigés accompagnent le %s du Centre VÉRITAS : ils en donnent les '
                'corrections, jamais les énoncés. Les exercices, les textes d\'auteur et les leçons sont dans '
                'l\'ouvrage — vendu en librairie ou accessible par abonnement ; la conduite de classe et les '
                'barèmes détaillés, dans le Guide pédagogique de l\'enseignant.</p>' % quoi)
    body.append("</div>")
    return h + "\n".join(body) + foot(depth=2)

def short_title(t):
    m = re.match(r"S[ée]quence\s+(\d+)", t, re.I) or re.match(r"S[ÉE]QUENCE\s+(\d+)", t)
    if m:
        return "Séq. %s" % m.group(1)
    m = re.match(r"Module\s+(\d+)", t, re.I)
    if m:
        return "Module %s" % m.group(1)
    if "diagnost" in t.lower():
        return "Diagnostic"
    if "BEPC" in t.upper():
        return "Vers le BEPC"
    if "CAP" in t.upper():
        return "Vers le CAP"
    if "démarche" in t.lower() or "DÉMARCHE" in t:
        return "Méthodes"
    if "sujet" in t.lower():
        return "Sujets d'examen"
    return t[:26]

def render_level_index(niv, pages, cahier=()):
    """Index d'un niveau. Deux ouvrages, donc deux axes : le cahier (par module) et
    le livret d'activités (par séquence). Le cahier passe en premier : c'est celui
    que l'élève a sous la main, et celui dont les corrigés n'existaient nulle part."""
    total = sum(p.n_items for p in pages) + sum(p.n_items for p in cahier)
    # Le 1er cycle découpe son cahier en modules, le 2nd cycle en séquences.
    decoupe = "module par module" if niv["cycle"] == "1er" else "séquence par séquence"
    exam = (" Préparation au %s." % niv["examen"]) if niv["examen"] else ""
    title = "Corrigés de français %s — %s exercices | VÉRITAS" % (niv["long"], num(total))
    desc = ("%s exercices de français corrigés pour la %s : le cahier %s et le livret "
            "d'activités, programme MINESEC.%s"
            % (num(total), niv["long"].replace("Classe de ", ""), decoupe, exam))
    url = "%s/corriges/%s/" % (SITE, niv["slug"])
    crumbs = [("Accueil", SITE), ("Corrigés des manuels", SITE + "/corriges/"), (niv["long"], url)]
    h = head(title, desc, url, depth=2, jsonld=jsonld_page(title, desc, url, niv["long"], crumbs))

    def cartes(lst):
        return "".join('<div class="card"><h3>%s</h3><p class="note">%s exercices corrigés</p>'
                       '<a class="dl" href="%s.html"><span>Voir les corrigés</span>'
                       '<span class="pill">%s</span></a></div>'
                       % (esc(p.titre), num(p.n_items), p.slug, short_title(p.titre)) for p in lst)

    ouvrages = ("le cahier « %s » et le livret d'activités" % niv["manuel"]) if cahier \
        else ("le livret d'activités « %s »" % niv.get("livret", niv["manuel"]))
    body = ['<header class="top"><span class="badge">Espace Manuels · Corrigés</span>'
            "<h1>Corrigés — %s</h1><p>%s exercices corrigés pour %s. Programme MINESEC.%s</p></header>"
            % (esc(niv["long"]), num(total), esc(ouvrages), esc(exam)),
            '<div class="wrap">',
            '<div class="intro">' + ico("i-target", "i lg") + ' <strong>La règle d\'or :</strong> on cherche d\'abord, on compare ensuite. '
            "Chaque corrigé est masqué tant que tu ne cliques pas — c'est fait exprès.</div>",
            '<p class="crumb"><a href="%s/">Accueil</a> › <a href="../">Corrigés des manuels</a> › %s</p>'
            % (SITE, esc(niv["long"]))]
    if cahier:
        body += ['<h2 class="sec">' + ico("i-notebook") + 'Le cahier de français, %s</h2>' % decoupe,
                 '<p class="note">Corrigés de « %s » — le cahier que tu utilises en classe, '
                 'des lectures méthodiques aux évaluations de fin de %s.</p>'
                 % (esc(niv["manuel"]), "module" if niv["cycle"] == "1er" else "séquence"),
                 '<div class="grid">%s</div>' % cartes(cahier)]
    body += ['<h2 class="sec">' + ico("i-book") + 'Le livret d\'activités, séquence par séquence</h2>',
             '<div class="grid">%s</div>' % cartes(pages),
            '<h2 class="sec">' + ico("i-compass") + 'Pour aller plus loin</h2>',
            '<div class="grid"><div class="card"><h3>Le programme de %s</h3>'
            '<p class="note">Notions, œuvres et examen du niveau.</p>'
            '<a class="dl" href="%s/niveaux/%s.html"><span>Voir le programme</span><span class="pill">MINESEC</span></a></div>'
            '<div class="card"><h3>Méthodes &amp; annales</h3>'
            '<p class="note">Résumé, commentaire, dissertation, dictée, épreuves corrigées.</p>'
            '<a class="dl" href="%s/ressources/"><span>Voir les méthodes</span><span class="pill">Fiches</span></a></div>'
            '<div class="card"><h3>Œuvres au programme</h3>'
            '<p class="note">Résumés, analyses et citations vérifiées.</p>'
            '<a class="dl" href="%s/oeuvres/"><span>Voir les œuvres</span><span class="pill">21 fiches</span></a></div></div>'
            % (esc(niv["label"]), SITE, NIVEAU_SEO[niv["slug"]], SITE, SITE),
            CTA, "</div>"]
    return h + "\n".join(body) + foot(depth=2)

def maj_manuels_html(levels):
    """Réécrit les cartes de `manuels.html` avec les compteurs réels.

    Ces nombres étaient saisis à la main : ils dérivaient à chaque régénération, et
    la page annonçait « corrigés des livrets et des bords » en ne pointant que sur
    les livrets. On les recalcule ici, à la source.
    """
    p = os.path.join(ROOT, "manuels.html")
    if not os.path.exists(p):
        return False
    src = open(p, encoding="utf-8").read()
    deb = src.find('<div class="grid" id="niveaux">')
    if deb < 0:
        return False
    fin = src.find("</div>\n", src.find('id="niveaux"'))
    fin = src.find("\n  </div>", deb)
    if fin < 0:
        return False
    cartes = []
    for niv, pages, cah in levels:
        tot = sum(x.n_items for x in pages) + sum(x.n_items for x in cah)
        ex = ' <small class="t">(%s)</small>' % niv["examen"] if niv["examen"] else ""
        quoi = "Livret — %d séquence%s" % (len(pages), "s" if len(pages) > 1 else "")
        if cah:
            quoi += " · Cahier — %d module%s" % (len(cah), "s" if len(cah) > 1 else "")
        elif niv["examen"]:
            quoi += " + sujets d'examen"
        cartes.append(
            '    <div class="card"><h3>%s%s</h3>\n'
            '      <a class="dl" href="corriges/%s/"><span>%sCorrigés · %s</span>'
            '<span class="pill">%s exos</span></a></div>'
            % (esc(niv["long"].replace("Classe de ", "Classe de ")), ex, niv["slug"],
               ico("i-book"), quoi, num(tot)))
    bloc = '<div class="grid" id="niveaux">\n' + "\n".join(cartes)
    open(p, "w", encoding="utf-8").write(src[:deb] + bloc + src[fin:])
    return True


def render_hub(levels, est):
    total = sum(sum(p.n_items for p in pages) + sum(p.n_items for p in cah)
                for _, pages, cah in levels)
    title = "Corrigés de français 6ᵉ à Terminale — %s exercices | VÉRITAS" % num(total)
    desc = ("%s exercices de français corrigés, de la 6ᵉ à la Terminale : cahiers et livrets "
            "du Centre VÉRITAS, programme MINESEC du Cameroun. Séries A, C, D et techniques." % num(total))
    url = SITE + "/corriges/"
    crumbs = [("Accueil", SITE), ("Corrigés des manuels", url)]
    h = head(title, desc, url, depth=1, jsonld=jsonld_page(title, desc, url, "Secondaire", crumbs))
    def group(items):
        out = []
        for niv, pages, cah in items:
            tot = sum(p.n_items for p in pages) + sum(p.n_items for p in cah)
            ex = ' <small class="note">(%s)</small>' % niv["examen"] if niv["examen"] else ""
            unite = "modules" if niv["cycle"] == "1er" else "séquences"
            detail = ("%d %s du cahier · %d séquences du livret" % (len(cah), unite, len(pages))) if cah \
                else ("%d séquences" % len(pages))
            out.append('<div class="card"><h3>%s%s</h3><p class="note">%s exercices corrigés · %s</p>'
                       '<a class="dl" href="%s/"><span>Voir les corrigés</span><span class="pill">%s exos</span></a></div>'
                       % (esc(niv["long"]), ex, num(tot), detail, niv["slug"], num(tot)))
        return "".join(out)
    c1 = group([x for x in levels if x[0]["cycle"] == "1er"])
    c2 = group([x for x in levels if x[0]["cycle"] == "2nd"])
    ce = "".join('<div class="card"><h3>%s</h3><p class="note">%s</p>'
                 '<a class="dl" href="%s"><span>Voir les corrigés</span><span class="pill">EN LIGNE</span></a></div>'
                 % (esc(t), esc(s), f) for f, t, s in est)
    body = ['<header class="top"><span class="badge">Centre VÉRITAS · Espace Manuels</span>'
            "<h1>Corrigés des cahiers de français — 6ᵉ à Terminale</h1>"
            "<p>%s exercices corrigés, séquence par séquence, pour les cahiers du Centre VÉRITAS. "
            "Programme MINESEC · Cameroun.</p></header>" % num(total),
            '<div class="wrap">',
            '<div class="intro">' + ico("i-book-open", "i lg") + ' <strong>C\'est la page annoncée dans ton cahier</strong> (encadré '
            + ico("i-globe") + ' Espace VÉRITAS '
            "et QR code de la couverture). Choisis ta classe, puis ta séquence. Les corrigés sont masqués "
            "tant que tu ne cliques pas : <strong>cherche d'abord, compare ensuite</strong>.</div>",
            '<p class="crumb"><a href="%s/">Accueil</a> › Corrigés des manuels</p>' % SITE,
            '<h2 class="sec">' + ico("i-book") + 'Premier cycle (6ᵉ → 3ᵉ)</h2><div class="grid">%s</div>' % c1,
            '<h2 class="sec">' + ico("i-book") + 'Second cycle — séries A (2ⁿᵈᵉ → Terminale)</h2>'
            '<div class="grid">%s</div>' % c2,
            '<h2 class="sec">' + ico("i-flask") + 'Séries scientifiques &amp; techniques (Cahiers EST)</h2>'
            '<p class="note">Corrigés des évaluations et des tâches d\'intégration : langue (méthode R.A.I.), '
            'résumés-modèles au décompte exact, plans d\'analyse et de dissertation.</p>'
            '<div class="grid">%s</div>' % ce,
            '<h2 class="sec">' + ico("i-compass") + 'Autres ressources VÉRITAS</h2>'
            '<div class="grid">'
            '<div class="card"><h3>Programmes par classe</h3><p class="note">Notions et œuvres, 6ᵉ → Tˡᵉ.</p>'
            '<a class="dl" href="%s/niveaux/"><span>Voir</span><span class="pill">7 pages</span></a></div>'
            '<div class="card"><h3>Méthodes &amp; examens</h3><p class="note">BEPC, Probatoire, BAC, GCE.</p>'
            '<a class="dl" href="%s/ressources/"><span>Voir</span><span class="pill">11 fiches</span></a></div>'
            '<div class="card"><h3>Œuvres au programme</h3><p class="note">Analyses et citations vérifiées.</p>'
            '<a class="dl" href="%s/oeuvres/"><span>Voir</span><span class="pill">21 œuvres</span></a></div>'
            '<div class="card"><h3>Espace Manuels</h3><p class="note">La page des manuels et de leurs ressources.</p>'
            '<a class="dl" href="%s/manuels.html"><span>Voir</span><span class="pill">Hub</span></a></div>'
            '</div>' % (SITE, SITE, SITE, SITE),
            CTA, "</div>"]
    return h + "\n".join(body) + foot(depth=1)

# ─────────────────────────────────────────────────────────────────── main
def main():
    os.makedirs(OUT, exist_ok=True)
    # La feuille de style et le JS sont communs à tout le site : /assets/veritas-pages.css
    # et /assets/veritas-ui.js (écrits à la main, versionnés). Rien à générer ici.

    levels, urls, report = [], [], []
    for niv in NIVEAUX:
        pages, stats = (extract_1er(niv) if niv["cycle"] == "1er" else extract_2nd(niv))
        cahier, stats_c = extract_cahier(niv)
        d = os.path.join(OUT, niv["slug"])
        os.makedirs(d, exist_ok=True)
        for p in pages:
            with open(os.path.join(d, p.slug + ".html"), "w", encoding="utf-8") as f:
                f.write(render_sequence_page(niv, p, pages, axe="livret"))
            urls.append("%s/corriges/%s/%s.html" % (SITE, niv["slug"], p.slug))
        for p in cahier:
            with open(os.path.join(d, p.slug + ".html"), "w", encoding="utf-8") as f:
                f.write(render_sequence_page(niv, p, cahier, axe="cahier"))
            urls.append("%s/corriges/%s/%s.html" % (SITE, niv["slug"], p.slug))
        with open(os.path.join(d, "index.html"), "w", encoding="utf-8") as f:
            f.write(render_level_index(niv, pages, cahier))
        urls.append("%s/corriges/%s/" % (SITE, niv["slug"]))
        levels.append((niv, pages, cahier))
        tot = sum(p.n_items for p in pages)
        totc = sum(p.n_items for p in cahier)
        report.append((niv["slug"], tot + totc, len(pages) + len(cahier), dict(stats)))
        print("  ✓ %-5s livret %2d pages/%4d exos · cahier %2d pages/%4d exos"
              % (niv["slug"], len(pages), tot, len(cahier), totc))

    est = [("est-2nde.html", "2ⁿᵈᵉ scientifique & technique", "Les Tribus de Capitoline · Le Tartuffe"),
           ("est-1ere.html", "1ʳᵉ scientifique & technique (Probatoire)", "Au cœur des ténèbres · Le Lion et la Perle"),
           ("est-tle.html", "Terminale scientifique & technique (BAC)", "Ngum a Jemea · Le Vieux Nègre et la Médaille")]
    with open(os.path.join(OUT, "index.html"), "w", encoding="utf-8") as f:
        f.write(render_hub(levels, est))
    print("  %s manuels.html (compteurs)" % ("✓" if maj_manuels_html(levels) else "!!"))
    urls.insert(0, SITE + "/corriges/")
    for f_, _, _ in est:
        urls.append("%s/corriges/%s" % (SITE, f_))

    with open(os.path.join(OUT, "sitemap-corriges.xml"), "w", encoding="utf-8") as f:
        f.write('<?xml version="1.0" encoding="UTF-8"?>\n'
                '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n')
        for u in dict.fromkeys(urls):
            prio = "0.9" if u.endswith("/corriges/") else ("0.8" if u.endswith("/") else "0.7")
            f.write("  <url><loc>%s</loc><lastmod>%s</lastmod>"
                    "<changefreq>monthly</changefreq><priority>%s</priority></url>\n" % (u, TODAY, prio))
        f.write("</urlset>\n")
    print("  ✓ hub + sitemap (%d URL)" % len(dict.fromkeys(urls)))
    print("\nTOTAL exercices publiés : %d" % sum(r[1] for r in report))

if __name__ == "__main__":
    main()
